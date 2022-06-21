from docx import Document 
from docx.shared import Cm 
 
from fake_contact import fake_contact 
 
from faker import Faker 
from faker.providers import lorem 
fake = Faker('es_ES') 
 
document = Document()    # Creación del documento 
 
seccion = document.seccions[-1]     # acceso a la sección 
seccion.left_margin = Cm(1) 
seccion.right_margin = Cm(1) 
seccion.top_margin = Cm(1) 
seccion.bottom_margin = Cm(1) 
 
# añadir un título 
document.add_heading('prueba de docx', 0) 
 
# añadir un párrafo 
p = document.add_paragraph( fake.text() ) 
 
# añadir otro párrafo
p = document.add_paragraph( 'Atributos de texto => ') 
 
# Añadir texto en el último párrafo
p.add_run('podemos añadir texto: ') 
p.add_run('en NEGRITA').bold = True 
p.add_run(' y texto: ') 
p.add_run('en CURSIVA.').italic = True 
 
data = [ ] 
for x in range(0,10): 
   data.append( fake_contact() ) 
 
## Añadir una tabla 
table = document.add_table(rows=1, cols=5) 
table.autofit = True 
hdr_cells = table.rows[0].cells 
hdr_cells[0].text = 'Nombre' 
hdr_cells[1].text = 'Feecha de nacimiento' 
hdr_cells[2].text = 'Dirección' 
hdr_cells[3].text = 'Teléfono' 
hdr_cells[4].text = 'Comentario' 
for n, d, a, t, c in data: 
    row_cells = table.add_row().cells 
    row_cells[0].text = n 
    row_cells[1].text = d.strftime("%d/%m/%Y") 
    row_cells[2].text = a 
    row_cells[3].text = t 
    row_cells[4].text = c 
 
# añadir un salto de página 
document.add_page_break() 
 
# Guardar el documento 
document.save('document1.docx')
